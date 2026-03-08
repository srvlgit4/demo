import os
import re
import shutil
import asyncio
import zipfile
import html
import gc
import threading
from docx import Document
from flask import Flask
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes

# ==========================================
# DUMMY WEB SERVER (KEEPS RENDER AWAKE)
# ==========================================
web_app = Flask(__name__)

@web_app.route('/')
def home():
    return "Master Novel Bot is running 24/7!"

def run_web_server():
    port = int(os.environ.get("PORT", 8080))
    web_app.run(host="0.0.0.0", port=port)

# ==========================================
# CONFIGURATION
# ==========================================
TOKEN = os.environ.get("BOT_TOKEN", "8436841638:AAFz0JFN8fXxHqy5eQGFDLXeCUwn0JLcF4w") 
GROUP_ID = int(os.environ.get("GROUP_ID", "-1003745983576"))
DEFAULT_DOCX_CHUNK = 50
DEFAULT_EPUB_CHUNK = 500

# State Management
document_queue = asyncio.Queue()
user_chunk_sizes = {}
pending_uploads = {}

pattern_num = r"^(?:vol(?:ume)?\s*\d+\s*)?(?:chapter|ch|c|अध्याय|चैप्टर|#|नॉवेलटैप|उपन्यासटैप|सी|पेज|पृष्ठ|000)?\s*(\d+)(?:[:\s-]|$)"

# ==========================================
# LOGIC 1: DOCX SPLITTER (WITH TOC BYPASS)
# ==========================================
def split_docx_logic(input_path, output_dir, chunk_size, output_format):
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    doc = Document(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    collector = []
    generated_files = []

    current_start = 1
    target_chapter = None
    first_chapter_found = False

    def save_chunk(lines, start, end, format_type):
        ext = ".txt" if format_type == "txt" else ".docx"
        part_name = f"{start}_to_{end}-{base_name}{ext}" if end not in ["End", "Full"] else f"{end}-{base_name}{ext}"
        if end == "End": part_name = f"{start}_to_End-{base_name}{ext}"
        part_path = os.path.join(output_dir, part_name)

        if format_type == "txt":
            with open(part_path, "w", encoding="utf-8") as f: f.write("\n\n".join(lines))
        else:
            new_doc = Document()
            for line in lines: new_doc.add_paragraph(line)
            new_doc.save(part_path)
            del new_doc 

        generated_files.append(part_path)

    lines = [para.text.strip() for para in doc.paragraphs]
    del doc 

    def is_toc_entry(index):
        lines_checked = 0
        for j in range(index + 1, len(lines)):
            if not lines[j]: continue
            lines_checked += 1
            if lines_checked > 8: break
            if re.match(pattern_num, lines[j], re.IGNORECASE): return True
        return False

    for i, text in enumerate(lines):
        if text and not first_chapter_found:
            match = re.match(pattern_num, text, re.IGNORECASE)
            if match and not is_toc_entry(i):
                detected_num = int(match.group(1))
                current_start = detected_num
                target_chapter = detected_num + chunk_size
                first_chapter_found = True

        is_boundary = False
        if text and first_chapter_found:
            match = re.match(pattern_num, text, re.IGNORECASE)
            if match and int(match.group(1)) == target_chapter:
                is_boundary = True

        if is_boundary:
            if collector: save_chunk(collector, current_start, target_chapter - 1, output_format)
            collector = [text]
            current_start = target_chapter
            target_chapter += chunk_size
        else:
            collector.append(text)

    if collector:
        end_marker = "End" if first_chapter_found else "Full"
        save_chunk(collector, current_start, end_marker, output_format)

    del lines
    gc.collect() 
    return generated_files

# ==========================================
# LOGIC 2: TXT SPLITTER (NEW - HIGH SPEED)
# ==========================================
def split_txt_logic(input_path, output_dir, chunk_size, output_format):
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    collector = []
    generated_files = []

    current_start = 1
    target_chapter = None
    first_chapter_found = False

    def save_chunk(lines, start, end, format_type):
        ext = ".txt" if format_type == "txt" else ".docx"
        part_name = f"{start}_to_{end}-{base_name}{ext}" if end not in ["End", "Full"] else f"{end}-{base_name}{ext}"
        if end == "End": part_name = f"{start}_to_End-{base_name}{ext}"
        part_path = os.path.join(output_dir, part_name)

        if format_type == "txt":
            with open(part_path, "w", encoding="utf-8") as f: f.write("\n".join(lines))
        else:
            new_doc = Document()
            for line in lines: new_doc.add_paragraph(line)
            new_doc.save(part_path)
            del new_doc 

        generated_files.append(part_path)

    # Read TXT extremely fast
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
        lines = [line.strip() for line in f if line.strip()]

    def is_toc_entry(index):
        lines_checked = 0
        for j in range(index + 1, len(lines)):
            if not lines[j]: continue
            lines_checked += 1
            if lines_checked > 8: break
            if re.match(pattern_num, lines[j], re.IGNORECASE): return True
        return False

    for i, text in enumerate(lines):
        if text and not first_chapter_found:
            match = re.match(pattern_num, text, re.IGNORECASE)
            if match and not is_toc_entry(i):
                detected_num = int(match.group(1))
                current_start = detected_num
                target_chapter = detected_num + chunk_size
                first_chapter_found = True

        is_boundary = False
        if text and first_chapter_found:
            match = re.match(pattern_num, text, re.IGNORECASE)
            if match and int(match.group(1)) == target_chapter:
                is_boundary = True

        if is_boundary:
            if collector: save_chunk(collector, current_start, target_chapter - 1, output_format)
            collector = [text]
            current_start = target_chapter
            target_chapter += chunk_size
        else:
            collector.append(text)

    if collector:
        end_marker = "End" if first_chapter_found else "Full"
        save_chunk(collector, current_start, end_marker, output_format)

    del lines
    gc.collect() 
    return generated_files

# ==========================================
# LOGIC 3: HIGH-SPEED EPUB CRACKER
# ==========================================
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

def fast_html_to_text(raw_html):
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', raw_html)
    text = re.sub(r'<(script|style|head)[^>]*>.*?</\1>', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?(p|div|h[1-6]|br|tr|li)[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = html.unescape(text)
    return [line.strip() for line in text.split('\n') if line.strip()]

def split_epub_logic(input_path, output_dir, chunk_size, output_format):
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    clean_name = re.sub(r'[^\w\-_]', '_', base_name)
    generated_files = []

    def save_epub_chunk(lines, count):
        ext = ".txt" if output_format == "txt" else ".docx"
        part_name = f"Part_{count}-{clean_name}{ext}"
        part_path = os.path.join(output_dir, part_name)

        if output_format == "txt":
            with open(part_path, "w", encoding="utf-8") as f: f.write("\n\n".join(lines))
        else:
            new_doc = Document()
            for line in lines: new_doc.add_paragraph(line)
            new_doc.save(part_path)
            del new_doc 

        generated_files.append(part_path)

    try:
        text_buffer = []
        chapter_count = 0
        chunk_count = 1

        with zipfile.ZipFile(input_path, 'r') as epub_zip:
            html_files = [f for f in epub_zip.namelist() if f.lower().endswith(('.html', '.xhtml', '.htm'))]
            html_files.sort(key=natural_sort_key)

            for file_name in html_files:
                try:
                    content = epub_zip.read(file_name).decode('utf-8', errors='ignore')
                    lines = fast_html_to_text(content)

                    if lines:
                        text_buffer.extend(lines)
                        text_buffer.append("---")
                        chapter_count += 1

                    if chapter_count >= chunk_size:
                        save_epub_chunk(text_buffer, chunk_count)
                        text_buffer = []
                        chapter_count = 0
                        chunk_count += 1
                        gc.collect() 

                except Exception as e:
                    print(f"Skipping bad EPUB section {file_name}: {e}")
                    continue

        if text_buffer:
            save_epub_chunk(text_buffer, chunk_count)

        gc.collect()
        return generated_files
    except Exception as e:
        print(f"Total EPUB Zip failure: {e}")
        return []

# ==========================================
# BACKGROUND WORKER (THE QUEUE PROCESSOR)
# ==========================================
async def queue_worker():
    while True:
        job = await document_queue.get()
        context, status_msg = job['context'], job['status_msg']
        input_path, output_dir = job['input_path'], job['output_dir']
        base_name, file_name = job['base_name'], job['file_name']

        try:
            loop = asyncio.get_event_loop()
            format_name = "TXT" if job['format'] == "txt" else "DOCX"

            if job['type'] == 'docx':
                await status_msg.edit_text(f"⚡ Processing Fast DOCX: `{file_name}` into chunks of {job['chunk_size']} as {format_name}...")
                files = await loop.run_in_executor(None, split_docx_logic, input_path, output_dir, job['chunk_size'], job['format'])
                err_msg = "⚠️ No chapters found. Is formatting correct?"
                intro_msg = f"📚 **{base_name}**\n👤 Uploaded by: {job['user_mention']}\n📄 Format: {format_name}"

            elif job['type'] == 'txt':
                await status_msg.edit_text(f"⚡ Processing Fast TXT: `{file_name}` into chunks of {job['chunk_size']} as {format_name}...")
                files = await loop.run_in_executor(None, split_txt_logic, input_path, output_dir, job['chunk_size'], job['format'])
                err_msg = "⚠️ No chapters found. Is formatting correct?"
                intro_msg = f"📚 **{base_name}**\n👤 Uploaded by: {job['user_mention']}\n📄 Format: {format_name}"

            elif job['type'] == 'epub':
                await status_msg.edit_text(f"⚡ High-Speed EPUB Extraction: `{file_name}` into chunks of {job['chunk_size']} as {format_name}...")
                files = await loop.run_in_executor(None, split_epub_logic, input_path, output_dir, job['chunk_size'], job['format'])
                err_msg = "⚠️ No readable text found. EPUB is heavily corrupted."
                intro_msg = f"📚 **{base_name}**\n👤 Uploaded by: {job['user_mention']}\n🧩 Split size: {job['chunk_size']} chapters\n📄 Format: {format_name}"

            if not files:
                await status_msg.edit_text(err_msg)
                continue

            thread_id = None
            try:
                topic = await context.bot.create_forum_topic(chat_id=GROUP_ID, name=base_name[:128])
                thread_id = topic.message_thread_id
                await status_msg.edit_text(f"✅ Created topic: **{base_name[:64]}**\n📤 Sending files...")
                await context.bot.send_message(chat_id=GROUP_ID, message_thread_id=thread_id, text=intro_msg)

                if job['type'] in ['epub', 'txt']:
                    with open(input_path, 'rb') as orig:
                        await context.bot.send_document(chat_id=GROUP_ID, message_thread_id=thread_id, document=orig, caption=f"📁 Original {job['type'].upper()} File")
            except Exception as e:
                await status_msg.edit_text(f"⚠️ Topic Error: Sending to main chat.")

            for f in files:
                with open(f, 'rb') as doc:
                    msg = await status_msg.reply_document(document=doc, filename=os.path.basename(f))
                    try:
                        forward_args = {"chat_id": GROUP_ID, "from_chat_id": msg.chat.id, "message_id": msg.message_id}
                        if thread_id: forward_args["message_thread_id"] = thread_id
                        await context.bot.forward_message(**forward_args)
                    except: pass

            await status_msg.reply_text("🎉 Done! Processing complete.")

        except Exception as e:
            await status_msg.edit_text(f"❌ Error: {e}")
        finally:
            if os.path.exists(job['temp_dir']): shutil.rmtree(job['temp_dir'])
            document_queue.task_done()
            gc.collect() 

# ==========================================
# BOT HANDLERS
# ==========================================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    name = update.effective_user.first_name
    await update.message.reply_text(
        f"👋 Hello {name}!\n\n"
        f"Send me a **.docx**, **.txt**, or **.epub** file.\n"
        f"`/set 500` - Change custom chunk size."
    )

async def set_chunk_size(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        new_size = int(context.args[0])
        user_chunk_sizes[update.effective_user.id] = new_size
        await update.message.reply_text(f"✅ Custom split size set to **{new_size}** chapters.")
    except: await update.message.reply_text("⚠️ Example: `/set 100`")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    file_name = doc.file_name.lower()
    msg_id = update.message.message_id

    if not (file_name.endswith('.docx') or file_name.endswith('.epub') or file_name.endswith('.txt')):
        await update.message.reply_text("❌ Only .docx, .txt, or .epub files allowed.")
        return

    pending_uploads[msg_id] = {
        'document': doc,
        'user_mention': f"@{update.effective_user.username}" if update.effective_user.username else update.effective_user.first_name,
        'user_id': update.effective_user.id
    }

    if file_name.endswith('.docx'):
        keyboard = [[InlineKeyboardButton("📄 DOCX", callback_data=f"docx|docx|{msg_id}"), InlineKeyboardButton("📝 TXT", callback_data=f"docx|txt|{msg_id}")]]
        await update.message.reply_text("DOCX detected. Save chunks as:", reply_markup=InlineKeyboardMarkup(keyboard))
    elif file_name.endswith('.txt'):
        keyboard = [[InlineKeyboardButton("📄 DOCX", callback_data=f"txt|docx|{msg_id}"), InlineKeyboardButton("📝 TXT", callback_data=f"txt|txt|{msg_id}")]]
        await update.message.reply_text("TXT detected. Save chunks as:", reply_markup=InlineKeyboardMarkup(keyboard))
    else:
        keyboard = [
            [InlineKeyboardButton(f"📄 DOCX ({DEFAULT_EPUB_CHUNK})", callback_data=f"epub|docx_def|{msg_id}"),
             InlineKeyboardButton(f"📝 TXT ({DEFAULT_EPUB_CHUNK})", callback_data=f"epub|txt_def|{msg_id}")],
            [InlineKeyboardButton("⚙️ DOCX (Custom /set)", callback_data=f"epub|docx_cust|{msg_id}"),
             InlineKeyboardButton("⚙️ TXT (Custom /set)", callback_data=f"epub|txt_cust|{msg_id}")]
        ]
        await update.message.reply_text("EPUB detected. Choose format and chunk size:", reply_markup=InlineKeyboardMarkup(keyboard))

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data.split("|")
    job_type, action, msg_id = data[0], data[1], int(data[2])

    if msg_id not in pending_uploads:
        await query.edit_message_text("❌ Session expired. Please upload again.")
        return

    job_info = pending_uploads.pop(msg_id)
    document = job_info['document']
    user_id = job_info['user_id']

    file_name = document.file_name
    temp_dir = f"temp_{user_id}_{msg_id}"
    input_path = os.path.join(temp_dir, file_name)
    os.makedirs(os.path.join(temp_dir, "output"), exist_ok=True)

    await query.edit_message_text(f"📥 Downloading `{file_name}`...")
    await (await document.get_file()).download_to_drive(input_path)

    job_data = {
        'type': job_type,
        'update': update, 'context': context, 'status_msg': query.message,
        'temp_dir': temp_dir, 'input_path': input_path, 'output_dir': os.path.join(temp_dir, "output"),
        'file_name': file_name, 'base_name': os.path.splitext(file_name)[0][:64].strip(),
        'user_mention': job_info['user_mention']
    }

    if job_type in ['docx', 'txt']:
        job_data['format'] = action
        job_data['chunk_size'] = user_chunk_sizes.get(user_id, DEFAULT_DOCX_CHUNK)
    else:
        format_choice = action.split('_')[0]
        size_type = action.split('_')[1]

        job_data['format'] = format_choice
        job_data['chunk_size'] = user_chunk_sizes.get(user_id, DEFAULT_EPUB_CHUNK) if size_type == "cust" else DEFAULT_EPUB_CHUNK

    await document_queue.put(job_data)

# ==========================================
# MAIN RUNNER
# ==========================================
async def start_background_tasks(app: Application):
    asyncio.create_task(queue_worker())

def main():
    print("🤖 Starting Web Server for Render Keep-Alive...")
    threading.Thread(target=run_web_server, daemon=True).start()

    print("🤖 Ultimate Fast Cracker Bot Initializing...")
    app = Application.builder().token(TOKEN).post_init(start_background_tasks).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("set", set_chunk_size))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(CallbackQueryHandler(button_callback))
    
    print("🚀 Master Bot is LIVE! (Speed & Memory Optimized)")
    app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
